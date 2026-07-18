import io
import pypdf
import docx

def extract_text_from_pdf(content: bytes) -> str:
    pdf_file = io.BytesIO(content)
    reader = pypdf.PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text.strip()

def extract_text_from_docx(content: bytes) -> str:
    docx_file = io.BytesIO(content)
    doc = docx.Document(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text += paragraph.text + "\n"
    # Extract text from tables if any
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text]
            if row_text:
                text += " | ".join(row_text) + "\n"
    return text.strip()

def extract_text(content: bytes, filename: str) -> str:
    ext = filename.split(".")[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(content)
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(content)
    else:
        raise ValueError(f"Unsupported file extension: .{ext}")
